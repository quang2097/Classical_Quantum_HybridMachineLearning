import torch
import pandas as pd
from pandas import DataFrame
from sklearn.model_selection import train_test_split
from sklearn.tree import DecisionTreeClassifier
from sklearn.ensemble import BaggingClassifier
from sklearn.metrics import accuracy_score, classification_report
from imblearn.under_sampling import RandomUnderSampler
from sklearn.preprocessing import LabelEncoder
from docx import Document
from docx.shared import Inches
from sklearn.metrics import confusion_matrix
import numpy as np

# --- New Imports for Plotting the Tree ---
from sklearn.tree import plot_tree
import matplotlib.pyplot as plt
import io

# NEW IMPORT: Seaborn for heatmap visualization
import seaborn as sns
#from google.colab import drive
#drive.mount('/content/drive')
import create_docx


def decision_tree(void, df:DataFrame) -> None:
    device = torch.device("cuda" if torch.cuda.is_available() else "cpu")
    le = LabelEncoder()

    # Convert categorical and object columns to numeric
    for col in df.select_dtypes(include=['object']).columns:
        if col == 'label':
            df[col] = le.fit_transform(df[col])
        else:
            df[col] = le.fit_transform(df[col])

    # --- Get the Target Names (CORRECT) ---
    # le.classes_ contains the array of original string labels in order of their integer encoding
    target_names = le.classes_

    X = df.drop(["src_ip","dst_ip","proto","service","conn_state","missed_bytes","src_pkts","src_ip_bytes","dst_pkts","dst_ip_bytes","dns_query","dns_qclass","dns_qtype","dns_rcode","dns_AA","dns_RD","dns_RA","dns_rejected","ssl_version","ssl_cipher","ssl_resumed","ssl_established","ssl_subject","ssl_issuer","http_trans_depth","http_method","http_uri","http_version","http_request_body_len","http_response_body_len","http_status_code","http_user_agent","http_orig_mime_types","http_resp_mime_types","weird_name","weird_addl","weird_notice","label","type"], axis=1)
    y = df["type"]

    # --- FIX: Comprehensive handling of non-finite, extreme values, and data types ---
    # Convert all columns that can be numeric to numeric, coercing errors
    for col in X.columns:
        X[col] = pd.to_numeric(X[col], errors='coerce')

    # Replace any infinity values with NaN
    X = X.replace([np.inf, -np.inf], np.nan)

    # Impute NaN values with the mean of their respective columns
    # Use numeric_only=True for mean to avoid errors with non-numeric cols if any slipped through
    X = X.fillna(X.mean(numeric_only=True))

    # Cap extreme outliers to prevent 'value too large' issues.
    # This prevents values from becoming np.inf if scikit-learn internally downcasts to float32.
    # Cap to the 99.9th percentile and 0.1st percentile
    for col in X.columns:
        if pd.api.types.is_numeric_dtype(X[col]):
            upper_bound = X[col].quantile(0.999)
            lower_bound = X[col].quantile(0.001)
            X[col] = np.clip(X[col], a_min=lower_bound, a_max=upper_bound)

    # Finally, ensure all feature columns are of float64 type for consistency and precision
    X = X.astype(np.float64)

    # Split the data into train and test sets first to prevent data leakage
    X_train, X_test, y_train, y_test = train_test_split(X, y, test_size=0.3, random_state=42)

    # Initialize RandomUnderSampler
    rus = RandomUnderSampler(random_state=42)

    # Apply undersampling only to the training set
    X_train_resampled, y_train_resampled = rus.fit_resample(X_train, y_train)

    ## A. Decision Tree
    dt_clf = DecisionTreeClassifier(random_state=42)
    dt_clf.fit(X_train_resampled, y_train_resampled)
    y_pred_decision = dt_clf.predict(X_test)

    ## B. Bagging Classifier
    base_clf = DecisionTreeClassifier(class_weight='balanced', random_state=42)
    bg_clf = BaggingClassifier(estimator=base_clf, n_estimators=100, max_features=0.8, random_state=42)
    bg_clf.fit(X_train_resampled, y_train_resampled)
    y_pred_random = bg_clf.predict(X_test)

    # 5. Evaluate
    print("--- Decision Tree Report (y_pred_decision) ---")
    dt_report_str = classification_report(y_test, y_pred_decision, target_names=target_names, zero_division=0)
    print(dt_report_str)

    print("\n--- Bagging Classifier Report (y_pred_random) ---")
    bg_report_str = classification_report(y_test, y_pred_random, target_names=target_names, zero_division=0)
    print(bg_report_str)

    dt_accuracy = accuracy_score(y_test, y_pred_decision)
    bg_accuracy = accuracy_score(y_test, y_pred_random)
    print("\nDecision Tree Accuracy: {:.2f}".format(dt_accuracy))
    print("Bagging Classifier Accuracy: {:.2f}".format(bg_accuracy))


    # --- FUNCTION: Create Decision Tree Plot (Optimized for DOCX) ---
    def generate_dt_plot_bytes(dt_clf, X, target_names):
        """Generates a limited-depth Decision Tree plot and returns it as bytes."""
        plt.figure(figsize=(25, 12))
        plot_tree(dt_clf,
            feature_names=X.columns.tolist(),
            class_names=target_names.tolist(),
            filled=True,
            rounded=True,
            fontsize=10,
            max_depth=3)
        plt.title("Decision Tree Visualization (Max Depth 3)", fontsize=16)

        buf = io.BytesIO()
        plt.savefig(buf, format='png', bbox_inches='tight', dpi=300)
        plt.close()
        buf.seek(0)
        return buf

    # --- NEW FUNCTION: Create Confusion Matrix Plot ---
    def plot_confusion_matrix_bytes(cm, class_names):
        """Generates a heatmap plot of the confusion matrix and returns it as bytes."""
        plt.figure(figsize=(10, 8))
        # Use Seaborn for a clean, colored heatmap
        sns.heatmap(cm,
            annot=True,
            fmt='d',
            cmap='Blues',
            xticklabels=class_names,
            yticklabels=class_names,
            cbar=True
        )

        plt.ylabel('Actual Label')
        plt.xlabel('Predicted Label')
        plt.title('Decision Tree Confusion Matrix')

        buf = io.BytesIO()
        # Save the figure to the buffer
        plt.savefig(buf, format='png', bbox_inches='tight', dpi=150)
        plt.close()
        buf.seek(0)
        return buf

    # Generate the plots before creating the report
    print("\n--- Generating Decision Tree Visualization (Max Depth 3) for DOCX ---")
    dt_plot_buffer = generate_dt_plot_bytes(dt_clf, X_train_resampled, target_names) # Use the correct target names

    cm = confusion_matrix(y_test, y_pred_decision)
    print("Decision Tree Confusion Matrix:\n", cm)

    print("--- Generating Confusion Matrix Image for DOCX ---")
    cm_plot_buffer = plot_confusion_matrix_bytes(cm, target_names) # Generate the image buffer

# --- FEATURE: Generate DOCX Report ---

    def create_and_save_report(dt_report_str, bg_report_str, bg_accuracy, dt_plot_buffer, dt_accuracy, cm_plot_buffer):
        """Generates a DOCX report and saves it to Google Drive."""
        document = Document()

        # --- Decision Tree Results ---
        document.add_heading('Decision Tree Classifier Results (Trained on Undersampled Data)', level=1)
        document.add_paragraph('Evaluation metrics for the Decision Tree Classifier:')
        document.add_paragraph(dt_report_str, 'No Spacing').style.font.name = 'Consolas'
        document.add_paragraph(f'Decision Tree overall accuracy: {dt_accuracy:.2f}.')

        # Add the Confusion Matrix IMAGE
        document.add_heading('Decision Tree Confusion Matrix (Image)', level=2)
        document.add_paragraph('The confusion matrix as a heatmap shows classification errors (Actual vs. Predicted):')
        document.add_picture(cm_plot_buffer, width=Inches(4.5))


        # --- Add the DT Plot to the Report ---
        document.add_heading('Decision Tree Visualization (Max Depth 3)', level=2)
        document.add_paragraph('A visualization of the top 3 levels of the Decision Tree.')
        document.add_picture(dt_plot_buffer, width=Inches(6.5))

        # --- Bagging Classifier Results ---
        document.add_heading('Bagging Classifier Results (Trained on Undersampled Data)', level=1)
        document.add_paragraph('Evaluation metrics for the Bagging Classifier:')
        document.add_paragraph(bg_report_str, 'No Spacing').style.font.name = 'Consolas'
        document.add_paragraph(f'Bagging Classifier overall accuracy: {bg_accuracy:.2f}.')

        # Add the Confusion Matrix IMAGE
        document.add_heading('Decision Tree Confusion Matrix (Image)', level=2)
        document.add_paragraph('The confusion matrix as a heatmap shows classification errors (Actual vs. Predicted):')
        document.add_picture(cm_plot_buffer, width=Inches(4.5))

        # Save the document to your Google Drive
        output_path = '/content/drive/MyDrive/Decision tree and bagging classifier ToN-IoT_Undersampled_with_CM_Image.docx'
        document.save(output_path)
        print(f"\nReport successfully saved to Google Drive at: {output_path}")

    # Run the function to create and save the report with the new confusion matrix image buffer
    create_docx.create_docx_decision_tree(dt_report_str, dt_plot_buffer, dt_accuracy, cm_plot_buffer, output_path)