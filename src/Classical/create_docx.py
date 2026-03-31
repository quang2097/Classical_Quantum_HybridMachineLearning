from docx import Document
from docx.shared import Inches

#create decision tree report docx
def create_docx_decision_tree(dt_report_str, dt_plot_buffer, dt_accuracy, cm_plot_buffer, output_path):
    document = Document()

    # --- Decision Tree Results ---
    document.add_heading('Decision Tree Classifier Results', level=1)
    document.add_paragraph('Evaluation metrics for the Decision Tree Classifier:')
    document.add_paragraph(dt_report_str, 'No Spacing').style.font.name = 'Consolas'
    document.add_paragraph(f'Decision Tree overall accuracy: {dt_accuracy:.2f}.')

    # Add the Confusion Matrix IMAGE
    document.add_heading('Decision Tree Confusion Matrix (Image)', level=2)
    document.add_paragraph('The confusion matrix as a heatmap shows classification errors (Actual vs. Predicted):')
    document.add_picture(cm_plot_buffer, width=Inches(4.5))

    # --- Add the DT Plot to the Report ---
    document.add_heading('Decision Tree Visualization (Max Depth 3)', level=2)
    document.add_paragraph('A visualization of the top 3 levels of the Decision Tree.')
    document.add_picture(dt_plot_buffer, width=Inches(6.5))

    # Save the document
    document.save(output_path)
    print(f"\nReport successfully saved to: {output_path}")

#create bagging classifier docx
def create_docx_bagging_classifier(bg_report_str, bg_accuracy, cm_plot_buffer, output_path):
    document = Document()

    # --- Bagging Classifier Results ---
    document.add_heading('Bagging Classifier Results (Trained on Undersampled Data)', level=1)
    document.add_paragraph('Evaluation metrics for the Bagging Classifier:')
    document.add_paragraph(bg_report_str, 'No Spacing').style.font.name = 'Consolas'
    document.add_paragraph(f'Bagging Classifier overall accuracy: {bg_accuracy:.2f}.')

    # Add the Confusion Matrix IMAGE
    document.add_heading('Decision Tree Confusion Matrix (Image)', level=2)
    document.add_paragraph('The confusion matrix as a heatmap shows classification errors (Actual vs. Predicted):')
    document.add_picture(cm_plot_buffer, width=Inches(4.5))

    # Save the document
    document.save(output_path)
    print(f"\nReport successfully saved to: {output_path}")

#create random forest docx
def create_docx_random_forest(accuracy, classification_report_str, feature_importance_list, importance_plot_buffer):
    """Generates a DOCX report for the Random Forest model and saves it to Google Drive."""
    document = Document()

    # Update heading to reflect the use of Undersampling
    document.add_heading('Random Forest Model Performance Report (Trained on Undersampled Data)', 0)

    # --- 1. Model Summary ---
    document.add_heading('1. Model Accuracy', level=1)
    document.add_paragraph(f'The trained Random Forest Classifier achieved an overall **Accuracy of {accuracy:.4f}** on the test dataset.')
    document.add_paragraph(f'The model was trained with 100 estimators and a maximum depth of 8, using **Random Under-Sampling (RUS)** on the training data to balance classes.')

    # --- 2. Classification Report ---
    document.add_heading('2. Detailed Classification Report', level=1)
    document.add_paragraph('The detailed per-class performance metrics (Precision, Recall, F1-Score) are as follows:')

    # Add the classification report as a pre-formatted paragraph to maintain alignment
    document.add_paragraph(classification_report_str, 'No Spacing').style.font.name = 'Consolas'

    # --- 3. Feature Importance ---
    document.add_heading('3. Feature Importance', level=1)

    # Add Feature Importance Plot
    document.add_heading('Top 10 Feature Importance Visualization', level=2)
    document.add_paragraph('A visualization of the top 10 most influential features according to the Gini importance metric:')
    document.add_picture(importance_plot_buffer, width=Inches(6.0))

    # Add Feature Importance List
    document.add_heading('All Feature Importance Scores', level=2)
    document.add_paragraph('The importance scores for all features (in descending order):')
    for item in feature_importance_list:
        document.add_paragraph(item, style='List Bullet')

    # Save the document to your Google Drive
    output_path = '/content/drive/MyDrive/Random_Forest_Report ToN-IoT_Undersampled.docx'
    document.save(output_path)
    print(f"\nReport successfully saved to Google Drive at: {output_path}")


#create lightBGM docx
def create_docx_lightBGM(auc_score, accuracy, classification_report_str, top_feature_importance_list, importance_plot_buffer, training_log_str, roc_plot_buffer, output_path):
    document = Document()

    # Update heading to reflect the use of RUS
    document.add_heading('LightGBM Model Performance Report (Trained on Undersampled Data)', 0)

    # --- 1. Model Summary ---
    document.add_heading('1. Overall Model Metrics', level=1)
    document.add_paragraph('The LightGBM model was trained using **Random Under-Sampling (RUS)** to mitigate class imbalance, and evaluated on the test set, yielding the following results:')
    document.add_paragraph(f'• Test **AUC Score** (Macro-Average): **{auc_score:.4f}**')
    document.add_paragraph(f'• Test **Accuracy**: **{accuracy:.4f}**')

    # --- 2. Training and Evaluation Log (Epoch Report) ---
    document.add_heading('2. Training Progress Log (Epoch Report)', level=1)
    document.add_paragraph('The log below shows the multi-logloss metric for the training (resampled data) and evaluation sets for each boosting round (up to the point of early stopping):')
    # Add the training log as a pre-formatted paragraph to maintain alignment
    document.add_paragraph(training_log_str, 'No Spacing').style.font.name = 'Consolas'

    # --- 3. Detailed Classification Report ---
    document.add_heading('3. Detailed Classification Report', level=1)
    document.add_paragraph('The detailed per-class performance metrics (Precision, Recall, F1-Score) are as follows:')

    # Add the classification report as a pre-formatted paragraph to maintain alignment
    document.add_paragraph(classification_report_str, 'No Spacing').style.font.name = 'Consolas'

    # --- 4. ROC Curve Analysis (NEW SECTION) ---
    document.add_heading('4. ROC Curve Analysis', level=1)
    document.add_paragraph('The Receiver Operating Characteristic (ROC) curve plots the True Positive Rate against the False Positive Rate for each class (One-vs-Rest):')
    document.add_picture(roc_plot_buffer, width=Inches(6.0))
    document.add_paragraph(f'The Macro-Average AUC Score is **{auc_score:.4f}**.')


    # --- 5. Feature Importance (Section Number changed from 4 to 5) ---
    document.add_heading('5. Feature Importance', level=1)

    # Add Feature Importance Plot
    document.add_heading('Top 10 Feature Importance Visualization', level=2)
    document.add_paragraph('A visualization of the top 10 most influential features:')
    document.add_picture(importance_plot_buffer, width=Inches(6.0))

    # Add Top 5 list
    document.add_heading('Top 5 Feature Scores', level=2)
    document.add_paragraph('The top 5 most influential features by score are:')
    for item in top_feature_importance_list:
        document.add_paragraph(item, style='List Bullet')

    # Save the document
    document.save(output_path)
    print(f"\nReport successfully saved to: {output_path}")

#create lightGBM docx
def create_docx_lightGBM():
    return

#create RNN docx
def create_docx_RNN():
    return

#create CNN docx
def create_docx_CNN():
    return

#create linearSVM docx
def create_docx_linearSVM():
    return

#create XGBoost docx
def create_docx_XGBoost():
    return

#create transformer docx
def create_docx_transformer():
    return

#create GNN docx
def create_docx_GNN():
    return 